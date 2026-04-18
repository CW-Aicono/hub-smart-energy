# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

BLUE = RGBColor(0x0B, 0x5E, 0xA8)
TEAL = RGBColor(0x0E, 0x9E, 0x9E)
GREY = RGBColor(0x55, 0x55, 0x55)

def shade(cell, hexcolor):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tc_pr.append(shd)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(6)

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = TEAL
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)

def para(text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.7 + 0.5*level)
    return p

def step(num, title, desc):
    p = doc.add_paragraph()
    r1 = p.add_run("Schritt " + str(num) + ": "); r1.bold = True; r1.font.color.rgb = BLUE; r1.font.size = Pt(13)
    r2 = p.add_run(title); r2.bold = True; r2.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(10)
    para(desc)

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F4F7')
    pPr.append(shd)
    return p

def tip(text, kind="info"):
    colors = {"info": "E8F4FD", "warn": "FFF4E5", "ok": "E7F6EC", "danger": "FDECEA"}
    icons  = {"info": "Tipp:", "warn": "Achtung:", "ok": "Gut zu wissen:", "danger": "Wichtig:"}
    tbl = doc.add_table(rows=1, cols=1); tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    shade(cell, colors[kind])
    p = cell.paragraphs[0]
    r = p.add_run(icons[kind] + " "); r.bold = True
    p.add_run(text)
    doc.add_paragraph()

# DECKBLATT
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("AICONO\nGateway-Worker"); r.bold = True; r.font.size = Pt(36); r.font.color.rgb = BLUE

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Installations-Anleitung"); r.font.size = Pt(20); r.font.color.rgb = TEAL

ver = doc.add_paragraph(); ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ver.add_run("Version 7 \u2014 Multi-Tenant"); r.font.size = Pt(14); r.font.color.rgb = GREY

doc.add_paragraph(); doc.add_paragraph()

intro_box = doc.add_table(rows=1, cols=1)
cell = intro_box.rows[0].cells[0]; shade(cell, "E8F4FD")
p = cell.paragraphs[0]
r = p.add_run("Was Sie hier finden\n"); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = BLUE
p.add_run(
    "\nEine Schritt-f\u00fcr-Schritt-Anleitung, mit der Sie GENAU ZWEI Container "
    "(einen f\u00fcr Live, einen f\u00fcr Staging) auf einem Hetzner-Server einrichten. "
    "Diese beiden Container bedienen ALLE Ihre Mandanten und ALLE Gateways "
    "automatisch \u2014 Sie m\u00fcssen nie wieder etwas anfassen, wenn ein neuer "
    "Mandant oder ein neues Gateway dazukommt."
)
doc.add_paragraph()
doc.add_page_break()

# INHALT
h1("Inhalt")
toc_items = [
    "1.  Das Konzept in 60 Sekunden",
    "2.  Was Sie brauchen (Einkaufsliste)",
    "3.  Hetzner-Server bestellen",
    "4.  Server vorbereiten (SSH, Updates, Sicherheit)",
    "5.  Docker installieren",
    "6.  Die zwei Schl\u00fcssel besorgen",
    "7.  Worker-Dateien hochladen",
    "8.  .env-Dateien ausf\u00fcllen",
    "9.  Container starten",
    "10. Funktioniert es? \u2014 Heartbeat pr\u00fcfen",
    "11. Wenn etwas nicht klappt (FAQ)",
    "12. Glossar \u2014 Fachbegriffe einfach erkl\u00e4rt",
]
for t in toc_items:
    p = doc.add_paragraph(t); p.paragraph_format.left_indent = Cm(0.5)
doc.add_page_break()

# 1
h1("1.  Das Konzept in 60 Sekunden")
para("Der Gateway-Worker ist ein kleines Programm, das rund um die Uhr l\u00e4uft "
     "und f\u00fcr Sie Daten von allen angeschlossenen Ger\u00e4ten (Loxone, Shelly, "
     "Tuya, ABB, Siemens, Home Assistant \u2026) abruft und in Ihre Cloud-Datenbank schreibt.")

h3("So funktioniert es bildlich:")
para("Stellen Sie sich den Worker wie einen Hausmeister vor, der f\u00fcr ALLE "
     "Ihre Liegenschaften gleichzeitig zust\u00e4ndig ist. Er macht alle 60 Sekunden "
     "die Runde und liest bei jedem einzelnen Ger\u00e4t den aktuellen Wert ab \u2014 egal "
     "wem das Geb\u00e4ude geh\u00f6rt. Die Cloud sagt ihm automatisch, welche Ger\u00e4te "
     "es gibt; Sie m\u00fcssen nichts mehr eintragen.")

tip("Neuer Mandant, neue Liegenschaft, neues Gateway? Sie machen NICHTS am "
    "Server. Sobald das Ger\u00e4t in der Cloud-Oberfl\u00e4che angelegt ist, wird es "
    "beim n\u00e4chsten 60-Sekunden-Lauf automatisch mit-versorgt.", "ok")

h3("Sie installieren genau 2 Container:")
bullet("Container 1: gateway-worker-live  \u2192  schreibt in die Live-Cloud (https://ems-pro.aicono.org)")
bullet("Container 2: gateway-worker-staging  \u2192  schreibt in die Test-Cloud (https://staging.aicono.org)")
para("Mehr nicht. Auch wenn Sie 200 Mandanten und 1.000 Geb\u00e4ude haben \u2014 "
     "es bleiben 2 Container.")
doc.add_page_break()

# 2
h1("2.  Was Sie brauchen (Einkaufsliste)")
h3("Was Sie KAUFEN m\u00fcssen:")
bullet("Einen Hetzner-Cloud-Account (kostenlos anlegen unter https://www.hetzner.com/cloud)")
bullet("Einen kleinen Server vom Typ \u201eCX22\u201c (2 CPU-Kerne, 4 GB RAM) \u2014 ca. 5 \u20ac/Monat")

h3("Was Sie auf Ihrem Computer brauchen:")
bullet("Einen normalen Webbrowser (Chrome, Edge, Firefox, Safari)")
bullet("Ein Terminal-Programm zum Verbinden mit dem Server:")
bullet("Windows: \u201ePowerShell\u201c (ist schon vorinstalliert)", level=1)
bullet("Mac: \u201eTerminal\u201c (ist schon vorinstalliert)", level=1)
bullet("Linux: jedes Terminal", level=1)

h3("Was Sie sich besorgen m\u00fcssen (machen wir sp\u00e4ter gemeinsam):")
bullet("Den Service-Role-Key aus Lovable Cloud (1 Schl\u00fcssel)")
bullet("Den Encryption-Key aus den Lovable-Secrets (1 Schl\u00fcssel)")

tip("Sie m\u00fcssen KEIN Programmierer sein. Wenn Sie einen Brief tippen k\u00f6nnen, "
    "k\u00f6nnen Sie diese Anleitung umsetzen. Alles, was zu tun ist: Befehle "
    "kopieren, einf\u00fcgen, Enter dr\u00fccken.", "info")
doc.add_page_break()

# 3
h1("3.  Hetzner-Server bestellen")
step(1, "Account anlegen",
     "Gehen Sie auf https://www.hetzner.com/cloud und klicken Sie auf "
     "\u201eJetzt registrieren\u201c. Folgen Sie der Anmeldung wie bei jedem anderen "
     "Online-Dienst (E-Mail, Passwort, Adresse, Zahlungsmethode).")
step(2, "Neues Projekt erstellen",
     "Nach dem Login klicken Sie oben links auf \u201e+ Neues Projekt\u201c. "
     "Geben Sie einen Namen ein, z. B. \u201eAICONO Gateway\u201c.")
step(3, "Server (Cloud Server) hinzuf\u00fcgen",
     "Im neuen Projekt klicken Sie auf \u201eServer hinzuf\u00fcgen\u201c. W\u00e4hlen Sie:")
bullet("Standort: N\u00fcrnberg oder Falkenstein (Deutschland)")
bullet("Image (Betriebssystem): Ubuntu 24.04")
bullet("Typ: CX22 (Shared vCPU, 2 Kerne, 4 GB RAM)")
bullet("Netzwerk: IPv4 + IPv6 (Standard, lassen)")
bullet("SSH-Schl\u00fcssel: Erstmal \u00fcberspringen \u2014 Hetzner schickt Ihnen das Passwort per E-Mail.")
bullet("Name: gateway-worker-aicono")
step(4, "Server bestellen",
     "Klicken Sie auf \u201eErstellen & kaufen\u201c. Nach 30 Sekunden ist der "
     "Server bereit. Notieren Sie sich:")
bullet("Die IP-Adresse (steht auf der Server-\u00dcbersichts-Seite, Format: 49.12.xxx.xxx)")
bullet("Das Root-Passwort (kommt per E-Mail von Hetzner)")
tip("Bewahren Sie die E-Mail mit dem Passwort sicher auf \u2014 Sie brauchen sie "
    "im n\u00e4chsten Schritt einmalig zum Einloggen.", "warn")
doc.add_page_break()

# 4
h1("4.  Server vorbereiten (SSH, Updates, Sicherheit)")
h3("4.1  Zum ersten Mal einloggen")
step(1, "Terminal \u00f6ffnen",
     "Windows: Windows-Taste, \u201ePowerShell\u201c tippen, Enter.\n"
     "Mac: Cmd+Leertaste, \u201eTerminal\u201c tippen, Enter.")
step(2, "Mit dem Server verbinden",
     "Tippen Sie folgenden Befehl (ersetzen Sie \u201eIHRE-IP\u201c durch die IP aus Hetzner):")
code("ssh root@IHRE-IP")
para("Beim ersten Mal fragt das Terminal:")
code('The authenticity of host ... can\'t be established. Are you sure you want to continue connecting (yes/no)?')
para("Tippen Sie \u201eyes\u201c und Enter. Dann werden Sie nach dem Passwort gefragt \u2014 das aus der Hetzner-E-Mail.")
tip("Beim Tippen des Passworts sehen Sie nichts auf dem Bildschirm \u2014 keine "
    "Punkte, keine Sterne, gar nichts. Das ist normal! Tippen Sie blind und Enter.", "info")

h3("4.2  Server aktualisieren")
para("Direkt nach dem Login geben Sie ein:")
code("apt update && apt upgrade -y")
para("Dauert 1\u20132 Minuten. Wenn ein farbiger Auswahl-Bildschirm erscheint, einfach Enter dr\u00fccken.")

h3("4.3  Firewall einschalten")
code("ufw allow OpenSSH\nufw --force enable")

h3("4.4  Automatische Sicherheitsupdates aktivieren")
code("apt install unattended-upgrades -y\ndpkg-reconfigure -plow unattended-upgrades")
para("Im Auswahl-Bildschirm: \u201eYes\u201c ausw\u00e4hlen, Enter.")
tip("Damit installiert Ihr Server zuk\u00fcnftig wichtige Sicherheitsupdates "
    "ganz von alleine. Sie m\u00fcssen sich um nichts mehr k\u00fcmmern.", "ok")
doc.add_page_break()

# 5
h1("5.  Docker installieren")
para("Docker ist das Programm, das die zwei Container starten und am Laufen halten wird. Installation in einem Befehl:")
code("curl -fsSL https://get.docker.com | sh")
para("Dauert ca. 1 Minute. Danach pr\u00fcfen Sie mit:")
code("docker --version")
para("Ausgabe sollte etwa so aussehen:")
code("Docker version 27.x.x, build xxxxxxx")
h3("Docker Compose installieren")
code("apt install docker-compose-plugin -y")
para("Pr\u00fcfen:")
code("docker compose version")
doc.add_page_break()

# 6
h1("6.  Die zwei Schl\u00fcssel besorgen")
para("Der Worker braucht zwei Schl\u00fcssel, damit er sicher mit der Cloud "
     "sprechen darf. Sie holen beide aus der Lovable-Oberfl\u00e4che.")

h3("6.1  SUPABASE_SERVICE_ROLE_KEY (Live)")
step(1, "Lovable Cloud \u00f6ffnen",
     "\u00d6ffnen Sie Ihr Lovable-Projekt im Browser und klicken Sie links in der Seitenleiste auf \u201eCloud\u201c.")
step(2, "Service-Role-Key anzeigen",
     "Klicken Sie auf \u201eAPI Keys\u201c (oder \u201eBackend > Settings > API\u201c). "
     "Sie sehen zwei Schl\u00fcssel: \u201eanon public\u201c und \u201eservice_role\u201c.")
step(3, "Service-Role-Key kopieren",
     "Klicken Sie beim \u201eservice_role\u201c-Schl\u00fcssel auf \u201eReveal\u201c und dann auf \u201eCopy\u201c. "
     "Speichern Sie ihn vorerst in einem sicheren Notizdokument.")
tip("Der service_role-Schl\u00fcssel ist EXTREM m\u00e4chtig \u2014 wer ihn hat, kann "
    "ALLE Daten lesen und \u00e4ndern. Geben Sie ihn niemals weiter, niemals in "
    "E-Mails, Chat oder Git!", "danger")

h3("6.2  SUPABASE_SERVICE_ROLE_KEY (Staging)")
para("Wiederholen Sie 6.1 \u2014 aber im Staging-Projekt von Lovable. Speichern "
     "Sie diesen Schl\u00fcssel separat in Ihrem Notizdokument unter \u201eSTAGING\u201c.")

h3("6.3  BRIGHTHUB_ENCRYPTION_KEY")
step(1, "Edge Function Secrets \u00f6ffnen",
     "Im Lovable-Projekt klicken Sie unter \u201eCloud\u201c auf \u201eEdge Functions\u201c \u2192 \u201eSecrets\u201c.")
step(2, "Schl\u00fcssel kopieren",
     "Suchen Sie den Eintrag \u201eBRIGHTHUB_ENCRYPTION_KEY\u201c. Kopieren Sie den Wert ins Notizdokument.")
para("Auch hier: separat f\u00fcr Live und Staging holen, falls beide Projekte verschiedene Schl\u00fcssel haben.")

h3("6.4  Ihre Notizen sollten jetzt so aussehen:")
code("LIVE:\n  SUPABASE_URL = https://xnveugycurplszevdxtw.supabase.co\n  SUPABASE_SERVICE_ROLE_KEY = eyJhbGc...sehr_lang...\n  BRIGHTHUB_ENCRYPTION_KEY = abc123...\n\nSTAGING:\n  SUPABASE_URL = https://staging-xxxxx.supabase.co\n  SUPABASE_SERVICE_ROLE_KEY = eyJhbGc...sehr_lang...\n  BRIGHTHUB_ENCRYPTION_KEY = xyz789...")
doc.add_page_break()

# 7
h1("7.  Worker-Dateien hochladen")
para("Das fertige Worker-Programm ist als ZIP-Datei verf\u00fcgbar (gateway-worker-src.zip). Diese laden wir auf den Server.")
h3("7.1  ZIP-Datei aus Lovable herunterladen")
para("Im Lovable-Projekt finden Sie die Datei unter docs/gateway-worker/gateway-worker-src.zip. "
     "Rechtsklick \u2192 \u201eSpeichern unter\u201c und auf Ihrem Computer ablegen.")
h3("7.2  ZIP auf den Server hochladen")
para("\u00d6ffnen Sie ein NEUES Terminal auf Ihrem Computer (das alte SSH-Fenster bleibt offen) und tippen Sie:")
code("scp gateway-worker-src.zip root@IHRE-IP:/root/")
para("Sie werden nach dem Server-Passwort gefragt. Eingeben, Enter.")
h3("7.3  ZIP auf dem Server entpacken")
para("Wechseln Sie zur\u00fcck ins SSH-Fenster und tippen Sie:")
code("cd /root\napt install unzip -y\nunzip gateway-worker-src.zip -d gateway-worker\ncd gateway-worker\nls")
para("Sie sollten jetzt diese Dateien sehen: Dockerfile, index.ts, package.json, tsconfig.json.")
doc.add_page_break()

# 8
h1("8.  .env-Dateien ausf\u00fcllen")
para("Jeder Container braucht seine eigene Konfigurations-Datei (.env). Wir erstellen zwei davon.")
h3("8.1  Live-Konfiguration")
code("nano /root/gateway-worker/.env.live")
para("Es \u00f6ffnet sich ein einfacher Texteditor. Tippen oder f\u00fcgen Sie ein:")
code("SUPABASE_URL=https://xnveugycurplszevdxtw.supabase.co\nSUPABASE_SERVICE_ROLE_KEY=HIER_DEN_LIVE_SERVICE_ROLE_KEY_EINFUEGEN\nBRIGHTHUB_ENCRYPTION_KEY=HIER_DEN_LIVE_ENCRYPTION_KEY_EINFUEGEN\nWORKER_ENV=live\nPOLL_INTERVAL_MS=60000\nHEARTBEAT_INTERVAL_MS=30000")
para("Speichern: Strg+O, dann Enter, dann Strg+X.")

h3("8.2  Staging-Konfiguration")
code("nano /root/gateway-worker/.env.staging")
code("SUPABASE_URL=https://staging-xxxxx.supabase.co\nSUPABASE_SERVICE_ROLE_KEY=HIER_DEN_STAGING_SERVICE_ROLE_KEY_EINFUEGEN\nBRIGHTHUB_ENCRYPTION_KEY=HIER_DEN_STAGING_ENCRYPTION_KEY_EINFUEGEN\nWORKER_ENV=staging\nPOLL_INTERVAL_MS=60000\nHEARTBEAT_INTERVAL_MS=30000")
para("Speichern wie zuvor (Strg+O, Enter, Strg+X).")

h3("8.3  docker-compose.yml erstellen")
code("nano /root/gateway-worker/docker-compose.yml")
code("services:\n  worker-live:\n    build: .\n    container_name: gateway-worker-live\n    restart: always\n    env_file: .env.live\n    logging:\n      driver: json-file\n      options:\n        max-size: \"10m\"\n        max-file: \"5\"\n\n  worker-staging:\n    build: .\n    container_name: gateway-worker-staging\n    restart: always\n    env_file: .env.staging\n    logging:\n      driver: json-file\n      options:\n        max-size: \"10m\"\n        max-file: \"5\"")
para("Speichern (Strg+O, Enter, Strg+X).")
tip("Achtung beim Einf\u00fcgen: Wenn Sie aus Ihrem Notizdokument kopieren, "
    "kann am Anfang oder Ende ein unsichtbares Leerzeichen mitkommen. "
    "Pr\u00fcfen Sie zur Sicherheit, dass nach \u201e=\u201c direkt der Schl\u00fcssel beginnt \u2014 kein Leerzeichen.", "warn")
doc.add_page_break()

# 9
h1("9.  Container starten")
para("Jetzt der Moment der Wahrheit \u2014 beide Container starten:")
code("cd /root/gateway-worker\ndocker compose up -d --build")
para("Das erste Mal dauert es 3\u20135 Minuten (Docker baut das Image). Danach sehen Sie:")
code("[+] Running 2/2\n  Container gateway-worker-live     Started\n  Container gateway-worker-staging  Started")
h3("Live-Logs ansehen")
code("docker logs -f gateway-worker-live")
para("Sie sollten Zeilen sehen wie \u201eDiscovery loop started\u201c, \u201eFound N integrations\u201c, "
     "\u201eHeartbeat sent\u201c. Mit Strg+C verlassen Sie die Log-Ansicht (der Container l\u00e4uft weiter).")
doc.add_page_break()

# 10
h1("10.  Funktioniert es? \u2014 Heartbeat pr\u00fcfen")
para("Es gibt zwei einfache Wege zu pr\u00fcfen, ob alles l\u00e4uft:")
h3("Weg 1: In der AICONO-Cloud")
bullet("Loggen Sie sich in https://ems-pro.aicono.org ein")
bullet("Gehen Sie als Super-Admin nach Einstellungen \u2192 Infrastruktur")
bullet("Sie sehen ein Widget \u201eGateway-Worker Status\u201c mit gr\u00fcnem Punkt")
bullet("Letzter Heartbeat sollte \u201evor wenigen Sekunden\u201c anzeigen")

h3("Weg 2: Direkt auf dem Server")
code("docker ps")
para("Beide Container m\u00fcssen mit Status \u201eUp\u201c angezeigt werden.")
code("docker logs gateway-worker-live --tail 20 | grep -i heartbeat")
para("Sie sollten regelm\u00e4\u00dfig \u201eHeartbeat sent\u201c sehen (alle 30 Sekunden).")
tip("Wenn der gr\u00fcne Punkt nach 5 Minuten noch immer da ist, ist die "
    "Installation komplett gegl\u00fcckt. Ab jetzt m\u00fcssen Sie an diesem Server "
    "NICHTS mehr tun \u2014 neue Mandanten und Gateways werden automatisch erkannt.", "ok")
doc.add_page_break()

# 11
h1("11.  Wenn etwas nicht klappt (FAQ)")
faqs = [
    ("Container startet nicht / stoppt sofort wieder",
     "Logs ansehen: docker logs gateway-worker-live\n"
     "H\u00e4ufigste Ursache: Tippfehler in der .env-Datei (z. B. Leerzeichen "
     "vor/nach dem Schl\u00fcssel). Datei mit nano \u00f6ffnen und pr\u00fcfen."),
    ("\u201eAuthentication failed\u201c / \u201eInvalid JWT\u201c in den Logs",
     "Der SUPABASE_SERVICE_ROLE_KEY ist falsch oder unvollst\u00e4ndig kopiert. "
     "Schl\u00fcssel aus Lovable Cloud nochmal frisch kopieren und .env aktualisieren."),
    ("Heartbeat-Widget bleibt rot",
     "1) docker ps pr\u00fcfen \u2014 laufen beide Container?\n"
     "2) docker logs gateway-worker-live ansehen \u2014 gibt es Fehler?\n"
     "3) Internet auf dem Server: ping -c 3 supabase.co"),
    ("Ich habe einen neuen Mandanten / ein neues Gateway angelegt \u2014 was tun?",
     "GAR NICHTS am Server. Beim n\u00e4chsten 60-Sekunden-Discovery-Lauf wird "
     "das Gateway automatisch mit-versorgt. Im Live-Log sehen Sie eine "
     "Zeile wie \u201eNew integration discovered: <name>\u201c."),
    ("Wie aktualisiere ich den Worker auf eine neue Version?",
     "1) Neue gateway-worker-src.zip per scp hochladen\n"
     "2) cd /root/gateway-worker && unzip -o ../gateway-worker-src.zip\n"
     "3) docker compose up -d --build"),
    ("Server-Passwort vergessen",
     "Im Hetzner-Cloud-Panel: Server ausw\u00e4hlen \u2192 \u201eRescue\u201c \u2192 neues Root-Passwort "
     "anfordern, Server neu starten."),
    ("Kann ich Worker-Live und Worker-Staging auf zwei verschiedene Server verteilen?",
     "Ja. Einfach diese Anleitung zweimal durchf\u00fchren \u2014 ein Server pro Container. "
     "Aber f\u00fcr Kostenersparnis reicht 1 CX22 f\u00fcr beide."),
]
for q, a in faqs:
    h3("Frage: " + q)
    para(a)
doc.add_page_break()

# 12
h1("12.  Glossar \u2014 Fachbegriffe einfach erkl\u00e4rt")
terms = [
    ("Container", "Wie eine kleine Schachtel, in der das Worker-Programm sicher und "
     "isoliert l\u00e4uft. Wenn der Container abst\u00fcrzt, startet Docker ihn automatisch neu."),
    ("Docker", "Das Programm, das diese Container verwaltet. Sie k\u00f6nnen Container "
     "starten, stoppen, Logs ansehen \u2014 alles mit kurzen Befehlen."),
    ("Docker Compose", "Eine Erweiterung von Docker, mit der man mehrere Container "
     "(in unserem Fall: live + staging) gemeinsam verwaltet."),
    ("SSH", "\u201eSecure Shell\u201c \u2014 die sichere Verbindung von Ihrem Computer zum Server."),
    ("Service-Role-Key", "Ein Master-Schl\u00fcssel f\u00fcr die Cloud-Datenbank. Damit "
     "kann der Worker Daten aller Mandanten lesen und schreiben. NIE \u00f6ffentlich machen!"),
    ("Encryption-Key", "Damit entschl\u00fcsselt der Worker die in der Cloud gespeicherten "
     "Gateway-Passw\u00f6rter (z. B. Loxone-Login)."),
    ("Heartbeat", "Ein \u201eLebenszeichen\u201c, das der Worker alle 30 Sekunden in die Cloud "
     "schickt. Bleibt der Heartbeat aus, wissen Sie sofort, dass etwas nicht stimmt."),
    ("Discovery-Loop", "Der 60-Sekunden-Rundgang, bei dem der Worker pr\u00fcft, ob neue "
     "Gateways oder Mandanten dazugekommen sind."),
    ("Tenant / Mandant", "Ein Kunde im System. Jeder Mandant hat seine eigenen "
     "Liegenschaften und Ger\u00e4te \u2014 alle Daten sind streng getrennt."),
    ("Liegenschaft / Location", "Ein konkretes Geb\u00e4ude (B\u00fcrogeb\u00e4ude, Schule \u2026)."),
    ("Gateway", "Ein Ger\u00e4t vor Ort, das Daten von Sensoren sammelt \u2014 Loxone-Miniserver, "
     "Shelly-Cloud-Account, Tuya-Hub etc."),
    ("Hetzner CX22", "Der kleine Cloud-Server, den wir bestellen. 2 CPU-Kerne, 4 GB RAM, "
     "ca. 5 \u20ac/Monat. Reicht f\u00fcr hunderte Mandanten."),
    ("nano", "Ein einfacher Texteditor auf dem Server. Bedienung: Pfeiltasten zum "
     "Bewegen, Strg+O zum Speichern, Strg+X zum Schlie\u00dfen."),
    ("Raspberry Pi", "Ein kleiner Computer f\u00fcr vor Ort. In dieser Architektur NICHT "
     "mehr n\u00f6tig \u2014 nur noch optional f\u00fcr Test- und Demo-Zwecke."),
]
for t, d in terms:
    p = doc.add_paragraph()
    r = p.add_run(t + ": "); r.bold = True; r.font.color.rgb = BLUE
    p.add_run(d)

doc.add_paragraph(); doc.add_paragraph()
footer = doc.add_paragraph(); footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("AICONO Gateway-Worker \u00b7 Version 7 \u00b7 Multi-Tenant\nBei Problemen: support@aicono.org")
r.font.size = Pt(10); r.font.color.rgb = GREY

doc.save('/mnt/documents/AICONO_Gateway_Worker_Installation_v7.docx')
print("OK saved")
